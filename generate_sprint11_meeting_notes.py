from __future__ import annotations

import json
import os
import re
import shutil
import sys
import time
from pathlib import Path
from typing import Any

from docx import Document
from dotenv import load_dotenv
from openai import OpenAI

DATASET_DIR = Path(__file__).resolve().parent
load_dotenv(DATASET_DIR / ".env")

MODEL = os.getenv("NUTRIVANA_OPENAI_MODEL", "openai/gpt-4o-mini")

# User referenced "sprint10_13_Meeting_Note.md"; in this repo the file is Sprint_10-13_Meeting_Notes.md.
SOURCE_MARKDOWN = DATASET_DIR / "Sprint_10-13_Meeting_Notes.md"
WORD_OUT_DIR = DATASET_DIR / "generated" / "word"
OUT_PATH = WORD_OUT_DIR / "sprint11_meeting_notes.docx"


class GenerationError(RuntimeError):
    pass


def _extract_sprint11_meeting_blueprint(full_md: str) -> str:
    start_tag = "# MEETING NOTE 2 — SPRINT 11"
    start = full_md.find(start_tag)
    if start == -1:
        raise GenerationError(f"Could not find {start_tag!r} in {SOURCE_MARKDOWN.name}")

    end_tag = "# MEETING NOTE 3 — SPRINT 12"
    end = full_md.find(end_tag, start)
    if end == -1:
        end = len(full_md)
    return full_md[start:end].strip()


def _system_prompt() -> str:
    """Same CRISPE system prompt style as the Sprint 1-9 meeting notes scripts."""
    sprint_number = 11
    return f"""ROLE: You are Shristi Sharmistha, CEO and PM at Nutrivana, taking live meeting notes during the mid-sprint weekly team sync recorded in the blueprint.

CONTEXT: Nutrivana is a nutrition tracking app for health-conscious Indians. Team is Shristi (PM), Arjun (CTO), Priya (Frontend), Kabir (Designer), Ananya (Marketing). This document is Sprint {sprint_number}'s weekly sync. Dates, sprint window, attendance, agenda, and specifics must follow the Sprint {sprint_number} meeting notes blueprint the user attaches.

INSTRUCTION: Read the Sprint {sprint_number} meeting notes blueprint carefully and write complete professional meeting notes. Every discussion point must capture what each person said, what was debated, and what decision was reached. Do not write summaries. Write like someone was in the room taking notes in real time.

PERFORMANCE:
- Sprint health check must be specific about which tickets are in progress and which are done
- Each key discussion must show who said what and what the outcome was
- Decisions must be specific and actionable
- Action items must have owner and due date
- Tone is direct startup energy not corporate

SPECIFICATION:
Format as a Word document with these sections:
- Header with date, sprint number, attendees
- Sprint Health Check
- Key Discussion 1 with full debate captured
- Key Discussion 2 with full debate captured
- Key Discussion 3 with full debate captured
- User and Market Updates
- Decisions Made as numbered list
- Action Items as table with Action, Owner, Due Date columns
- Next Week Focus"""


def _client() -> OpenAI:
    key = os.getenv("OPENAI_API_KEY", "").strip()
    base_url = os.getenv("OPENAI_BASE_URL", "").strip() or None
    if not key or key == "sk-or-v1-your-key-here":
        raise GenerationError(
            "OpenAI API key missing. Set OPENAI_API_KEY in nutrivana-dataset/.env (use python-dotenv)."
        )
    kwargs: dict[str, Any] = {"api_key": key}
    if base_url:
        kwargs["base_url"] = base_url
    return OpenAI(**kwargs)


def _validate_meeting_spec(data: dict[str, Any]) -> None:
    title = str(data.get("document_title", "")).strip()
    if not title:
        raise ValueError("document_title must be non-empty")
    for k in ("meeting_date", "sprint_line", "attendees"):
        if not str(data.get(k, "")).strip():
            raise ValueError(f"{k} must be non-empty string")
    if not str(data.get("sprint_health_check", "")).strip():
        raise ValueError("sprint_health_check must be non-empty")
    if not str(data.get("user_and_market_updates", "")).strip():
        raise ValueError("user_and_market_updates must be non-empty")
    if not str(data.get("next_week_focus", "")).strip():
        raise ValueError("next_week_focus must be non-empty")

    for n in (1, 2, 3):
        kd = data.get(f"key_discussion_{n}")
        if not isinstance(kd, dict):
            raise ValueError(f"key_discussion_{n} must be an object")
        if not str(kd.get("title", "")).strip():
            raise ValueError(f"key_discussion_{n}.title must be non-empty")
        if not str(kd.get("notes", "")).strip():
            raise ValueError(f"key_discussion_{n}.notes must be non-empty")

    decisions = data.get("decisions")
    if not isinstance(decisions, list) or len(decisions) < 1:
        raise ValueError("decisions must be a non-empty array of strings")
    for i, d in enumerate(decisions):
        if not str(d).strip():
            raise ValueError(f"decisions[{i}] must be non-empty")

    actions = data.get("action_items")
    if not isinstance(actions, list) or len(actions) < 1:
        raise ValueError("action_items must be a non-empty array")
    for i, a in enumerate(actions):
        if not isinstance(a, dict):
            raise ValueError(f"action_items[{i}] must be an object")
        for k in ("action", "owner", "due_date"):
            if not str(a.get(k, "")).strip():
                raise ValueError(f"action_items[{i}] needs non-empty {k}")


def _call_meeting_json(client: OpenAI, *, blueprint: str) -> dict[str, Any]:
    user = (
        "Return ONLY valid JSON (no markdown, no code fences) for the Sprint 11 weekly meeting notes Word document.\n"
        "Follow the SYSTEM prompt: live-capture dialogue, ticket-level sprint health, no generic summaries.\n"
        "Key discussion topics must align to the blueprint: (1) 38% drop-off confirmed, (2) custom food 2.3x retention and first-3-sessions insight, (3) March interview notes connection.\n\n"
        "JSON schema:\n"
        "{\n"
        '  "document_title": "string (match blueprint document name tone)",\n'
        '  "meeting_date": "string",\n'
        '  "sprint_line": "string (should read Sprint 11 (May 27–June 9, 2025))",\n'
        '  "attendees": "comma-separated names",\n'
        '  "sprint_health_check": "string multi-paragraph narrative",\n'
        '  "key_discussion_1": {"title": "38% drop-off confirmed", "notes": "long live notes with who said what"},\n'
        '  "key_discussion_2": {"title": "Custom food 2.3x retention + first-3-sessions insight", "notes": "..."},\n'
        '  "key_discussion_3": {"title": "March interview notes connection", "notes": "..."},\n'
        '  "user_and_market_updates": "string multi-paragraph",\n'
        '  "decisions": ["string decision 1", "..."],\n'
        '  "action_items": [{"action":"...","owner":"...","due_date":"..."}, ...],\n'
        '  "next_week_focus": "string multi-paragraph or bullets as prose"\n'
        "}\n\n"
        "Blueprint:\n-----\n"
        f"{blueprint}\n-----\n"
    )

    messages: list[dict[str, str]] = [
        {"role": "system", "content": _system_prompt()},
        {"role": "user", "content": user},
    ]

    max_attempts = 5  # 1 initial + 4 retries
    last_err: Exception | None = None
    for attempt in range(max_attempts):
        try:
            resp = client.chat.completions.create(
                model=MODEL,
                messages=messages,
                temperature=0.35,
                max_tokens=12000,
                response_format={"type": "json_object"},
            )
            text = (resp.choices[0].message.content or "").strip()
            data = json.loads(text)
            _validate_meeting_spec(data)
            return data
        except Exception as e:  # noqa: BLE001
            last_err = e
            if attempt == max_attempts - 1:
                break
            messages.append(
                {
                    "role": "user",
                    "content": (
                        "Your previous JSON was invalid or incomplete. Return ONLY valid JSON matching the schema. "
                        f"Error: {e}"
                    ),
                }
            )
            time.sleep(min(2.0, 0.5 * (attempt + 1)))
    raise GenerationError(f"Failed after retries: {last_err}")


def _add_body_paragraphs(doc: Document, text: str) -> None:
    blocks = [b.strip() for b in re.split(r"\n\s*\n", text.strip()) if b.strip()]
    for b in blocks:
        for line in b.split("\n"):
            line = line.strip()
            if line:
                doc.add_paragraph(line)


def _write_docx(spec: dict[str, Any], out_path: Path) -> Path:
    doc = Document()
    doc.add_heading(str(spec["document_title"]).strip(), level=0)

    doc.add_paragraph(f"Date: {str(spec['meeting_date']).strip()}")
    doc.add_paragraph(f"Sprint: {str(spec['sprint_line']).strip()}")
    doc.add_paragraph(f"Attendees: {str(spec['attendees']).strip()}")

    doc.add_heading("Sprint Health Check", level=1)
    _add_body_paragraphs(doc, str(spec["sprint_health_check"]))

    for n in (1, 2, 3):
        kd = spec[f"key_discussion_{n}"]
        title = str(kd["title"]).strip()
        doc.add_heading(f"Key Discussion {n} — {title}", level=1)
        _add_body_paragraphs(doc, str(kd["notes"]))

    doc.add_heading("User and Market Updates", level=1)
    _add_body_paragraphs(doc, str(spec["user_and_market_updates"]))

    doc.add_heading("Decisions Made", level=1)
    for i, d in enumerate(spec["decisions"], start=1):
        doc.add_paragraph(f"{i}. {str(d).strip()}")

    doc.add_heading("Action Items", level=1)
    headers = ["Action", "Owner", "Due Date"]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.cell(0, i).text = h
    for a in spec["action_items"]:
        row = table.add_row().cells
        row[0].text = str(a["action"]).strip()
        row[1].text = str(a["owner"]).strip()
        row[2].text = str(a["due_date"]).strip()

    doc.add_heading("Next Week Focus", level=1)
    _add_body_paragraphs(doc, str(spec["next_week_focus"]))

    out_path.parent.mkdir(parents=True, exist_ok=True)
    tmp = out_path.with_suffix(".tmp.docx")
    doc.save(tmp)
    try:
        os.replace(tmp, out_path)
        return out_path
    except PermissionError:
        fb = out_path.with_name("sprint11_meeting_notes_ready.docx")
        shutil.move(str(tmp), str(fb))
        return fb


def main() -> int:
    md = SOURCE_MARKDOWN.read_text(encoding="utf-8")
    blueprint = _extract_sprint11_meeting_blueprint(md)

    print(f"Model: {MODEL}")
    print("Calling API for Sprint 11 meeting notes...")
    client = _client()
    spec = _call_meeting_json(client, blueprint=blueprint)
    saved = _write_docx(spec, OUT_PATH)
    if saved != OUT_PATH:
        print(
            f"Saved to {saved} (could not overwrite {OUT_PATH} — file may be open in Word). "
            "Close sprint11_meeting_notes.docx and rename/copy the new file, or rerun."
        )
    else:
        print(f"Wrote {saved}")
    return 0


if __name__ == "__main__":
    try:
        raise SystemExit(main())
    except GenerationError as e:
        print(str(e), file=sys.stderr)
        raise SystemExit(1)

