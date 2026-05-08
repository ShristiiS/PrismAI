from __future__ import annotations

import json
import os
import shutil
import sys
import time
from pathlib import Path
from typing import Any

from docx import Document
from dotenv import load_dotenv
from openai import OpenAI

DATASET_DIR = Path(__file__).resolve().parent
load_dotenv(DATASET_DIR / ".env")

MODEL = os.getenv("NUTRIVANA_OPENAI_MODEL", "openai/gpt-4o-mini")

# User referenced sprint10_13_planning_retro_definitions.md; in this repo the source file is Sprint_10-13_Planning.md.
SOURCE_MARKDOWN = DATASET_DIR / "Sprint_10-13_Planning.md"
OUT_PATH = DATASET_DIR / "generated" / "word" / "sprint13_planning.docx"

SYSTEM_PROMPT = """ROLE: You are Shristi Sharmistha, CEO and Product Manager at Nutrivana. You have been running agile sprints for 6 months. You write sprint planning documents that your engineering team actually reads and follows.

CONTEXT: Nutrivana is a nutrition tracking app for health-conscious Indians. Tech stack is Next.js, FastAPI, Supabase, PostgreSQL. The team is Shristi (PM), Arjun Mehta (CTO), Priya Nair (Frontend Engineer), Kabir Sharma (Designer), Ananya Iyer (Marketing). The sprint number, calendar dates, and sprint theme are exactly as stated in the planning blueprint the user attaches — use that as source of truth.

INSTRUCTION: Read the sprint planning blueprint carefully and write a complete professional sprint planning document. Do not copy the blueprint word for word. Use it as your source of truth to write a natural professional document that a real startup PM would write. Every section must feel like it was written by a human who deeply understands the product and the team.

PERFORMANCE:
- Team capacity table must show each person with their available days and what they are focused on
- Sprint backlog table must explain WHY each ticket is in this sprint not just what it is
- Risks section must feel real — specific risks with specific mitigations not generic statements
- Definition of done must be concrete and testable
- Avoid corporate speak like 'leverage synergies' or 'ensure alignment'
- Avoid vague statements like 'complete the necessary tasks'
- Every sentence must add information a team member would actually need

SPRINT GOAL RULE:
The sprint goal must be written exactly like how a real startup PM writes it. It must:
- Be 1-2 sentences maximum
- Say specifically what is being built or achieved this sprint
- Say what is NOT being done if that is important context
- Use direct language — no corporate words like 'leverage' or 'enable seamless'
- Sound like something Shristi would say in a standup

EXAMPLE of good sprint goal for an infrastructure sprint:
'Establish the complete technical foundation before any feature development begins. No feature code ships this sprint — only database schema, authentication, and research.'

EXAMPLE of good sprint goal for a feature sprint:
'Complete the full food logging flow so a user can search food, adjust portions, add to their diary, and see nutrient totals update in real time. By end of sprint the product is demoable for the first time.'

EXAMPLE of bad sprint goal:
'Complete the necessary infrastructure tasks to enable future feature development and ensure the team is aligned.'

SPECIFICATION:
- Format as a properly structured Word document with clear headings
- Sprint goal: follow SPRINT GOAL RULE; render as one bold block (1-2 sentences) immediately under the title/subtitle
- Team Capacity: table with columns Person, Role, Available Days, Focus This Sprint
- Sprint Backlog: table with columns Ticket ID, Title, Assignee, Story Points, Priority, Why This Sprint
- Success Criteria: bullet points starting with a measurable outcome
- Dependencies: what must be true before this sprint can succeed
- Risks: bullet points with Risk and Mitigation for each
- Definition of Done: bullet points that are specific and testable
- Tone: direct, startup energy, no fluff"""


class GenerationError(RuntimeError):
    pass


def _extract_sprint13_planning_blueprint(full_md: str) -> str:
    start_tag = "# SPRINT 13 PLANNING DOCUMENT"
    end_tag = "# SPRINT 13 RETROSPECTIVE"
    start = full_md.find(start_tag)
    if start == -1:
        raise GenerationError(f"Could not find {start_tag!r} in {SOURCE_MARKDOWN.name}")
    end = full_md.find(end_tag, start)
    if end == -1:
        raise GenerationError(f"Could not find {end_tag!r} after Sprint 13 planning section")
    return full_md[start:end].strip()


def _client() -> OpenAI:
    key = os.getenv("OPENAI_API_KEY", "").strip()
    base_url = os.getenv("OPENAI_BASE_URL", "").strip() or None
    if not key or key == "sk-or-v1-your-key-here":
        raise GenerationError(
            "OpenAI API key missing. Set OPENAI_API_KEY in nutrivana-dataset/.env (use python-dotenv)."
        )
    kwargs: dict[str, Any] = {"api_key": key}
    if base_url:
        kwargs["base_url"] = base_url
    return OpenAI(**kwargs)


def _validate_planning_spec(data: dict[str, Any]) -> None:
    title = str(data.get("title", "")).strip()
    goal = str(data.get("sprint_goal", "")).strip()
    if not title:
        raise ValueError("Missing non-empty title")
    if not goal:
        raise ValueError("Missing non-empty sprint_goal")
    tables = data.get("tables")
    if not isinstance(tables, list) or len(tables) < 2:
        raise ValueError("Need tables array with Team Capacity and Sprint Backlog")

    tc = tables[0]
    bl = tables[1]
    tc_h = [str(x).strip() for x in (tc.get("headers") or [])]
    bl_h = [str(x).strip() for x in (bl.get("headers") or [])]
    exp_tc = ["Person", "Role", "Available Days", "Focus This Sprint"]
    exp_bl = ["Ticket ID", "Title", "Assignee", "Story Points", "Priority", "Why This Sprint"]
    if tc_h != exp_tc:
        raise ValueError(f"Team capacity headers must be {exp_tc}, got {tc_h}")
    if bl_h != exp_bl:
        raise ValueError(f"Sprint backlog headers must be {exp_bl}, got {bl_h}")

    if not (tc.get("rows") or []):
        raise ValueError("Team capacity table needs at least one row")
    if not (bl.get("rows") or []):
        raise ValueError("Sprint backlog table needs at least one row")

    for key in ("success_criteria", "dependencies", "definition_of_done"):
        arr = data.get(key) or []
        if not isinstance(arr, list) or not arr:
            raise ValueError(f"{key} must be a non-empty array of strings")

    risks = data.get("risks") or []
    if not isinstance(risks, list) or not risks:
        raise ValueError("risks must be a non-empty array of objects with risk and mitigation")
    for i, item in enumerate(risks):
        if not isinstance(item, dict):
            raise ValueError(f"risks[{i}] must be an object")
        if not str(item.get("risk", "")).strip() or not str(item.get("mitigation", "")).strip():
            raise ValueError(f"risks[{i}] needs non-empty risk and mitigation")


def _call_planning_json(client: OpenAI, *, user: str, max_retries: int = 4) -> dict[str, Any]:
    messages: list[dict[str, str]] = [
        {"role": "system", "content": SYSTEM_PROMPT},
        {"role": "user", "content": user},
    ]
    last_err: Exception | None = None
    for attempt in range(max_retries + 1):
        try:
            resp = client.chat.completions.create(
                model=MODEL,
                messages=messages,
                temperature=0.35,
                response_format={"type": "json_object"},
            )
            text = (resp.choices[0].message.content or "").strip()
            data = json.loads(text)
            _validate_planning_spec(data)
            return data
        except Exception as e:  # noqa: BLE001
            last_err = e
            messages.append(
                {
                    "role": "user",
                    "content": (
                        "Your previous reply was not usable. Return ONLY valid JSON matching the schema. "
                        f"Error: {e}"
                    ),
                }
            )
            time.sleep(min(2.0, 0.5 * (attempt + 1)))
    raise GenerationError(f"Failed after retries: {last_err}")


def _add_bullet_list(doc: Document, items: list[str]) -> None:
    for raw in items:
        t = str(raw).strip()
        if not t:
            continue
        doc.add_paragraph(t, style="List Bullet")


def _write_docx(data: dict[str, Any], out_path: Path) -> Path:
    doc = Document()
    doc.add_heading(str(data["title"]).strip(), level=0)
    sub = data.get("subtitle")
    if sub and str(sub).strip():
        p = doc.add_paragraph(str(sub).strip())
        for r in p.runs:
            r.italic = True

    gp = doc.add_paragraph()
    gr = gp.add_run(str(data["sprint_goal"]).strip())
    gr.bold = True

    tables = data.get("tables") or []
    for i, tbl in enumerate(tables[:2]):
        heading = "Team Capacity" if i == 0 else "Sprint Backlog"
        doc.add_heading(heading, level=1)
        headers = [str(h) for h in (tbl.get("headers") or [])]
        rows = tbl.get("rows") or []
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        for j, h in enumerate(headers):
            table.cell(0, j).text = h
        for r in rows:
            cells = table.add_row().cells
            for j, v in enumerate(r):
                cells[j].text = "" if v is None else str(v)

    doc.add_heading("Success Criteria", level=1)
    _add_bullet_list(doc, list(data.get("success_criteria") or []))

    doc.add_heading("Dependencies", level=1)
    _add_bullet_list(doc, list(data.get("dependencies") or []))

    doc.add_heading("Risks", level=1)
    for item in data.get("risks") or []:
        if not isinstance(item, dict):
            continue
        risk = str(item.get("risk", "")).strip()
        mit = str(item.get("mitigation", "")).strip()
        doc.add_paragraph(f"Risk: {risk} Mitigation: {mit}", style="List Bullet")

    doc.add_heading("Definition of Done", level=1)
    _add_bullet_list(doc, list(data.get("definition_of_done") or []))

    out_path.parent.mkdir(parents=True, exist_ok=True)
    tmp = out_path.with_suffix(".tmp.docx")
    doc.save(tmp)
    try:
        os.replace(tmp, out_path)
        return out_path
    except PermissionError:
        fb = out_path.with_name("sprint13_planning_ready.docx")
        shutil.move(str(tmp), str(fb))
        return fb


def main() -> int:
    full_md = SOURCE_MARKDOWN.read_text(encoding="utf-8")
    blueprint = _extract_sprint13_planning_blueprint(full_md)

    user = (
        "Return ONLY valid JSON (no markdown, no code fences) for a Sprint 13 planning Word document. "
        "Follow the SYSTEM prompt structure and tone. Use the blueprint as source of truth for names, "
        "dates, tickets, story points, and priorities — but rewrite in your own professional words.\n\n"
        "JSON schema (exact keys):\n"
        "{\n"
        '  "title": "string (e.g. Sprint 13 Planning — Nutrivana)",\n'
        '  "subtitle": "string optional (e.g. sprint dates / doc date)",\n'
        '  "sprint_goal": "string, 1-2 sentences following SPRINT GOAL RULE in the system prompt",\n'
        '  "tables": [\n'
        "    {\n"
        '      "headers": ["Person", "Role", "Available Days", "Focus This Sprint"],\n'
        '      "rows": [["...", "...", "...", "..."]]\n'
        "    },\n"
        "    {\n"
        '      "headers": ["Ticket ID", "Title", "Assignee", "Story Points", "Priority", "Why This Sprint"],\n'
        '      "rows": [["...", "...", "...", "...", "...", "..."]]\n'
        "    }\n"
        "  ],\n"
        '  "success_criteria": ["string measurable bullet", "..."],\n'
        '  "dependencies": ["string", "..."],\n'
        '  "risks": [{"risk": "string", "mitigation": "string"}, ...],\n'
        '  "definition_of_done": ["string testable bullet", "..."]\n'
        "}\n\n"
        "Blueprint:\n-----\n"
        f"{blueprint}\n"
        "-----\n"
    )

    print(f"Model: {MODEL}")
    print("Calling API for Sprint 13 planning document...")
    client = _client()
    spec = _call_planning_json(client, user=user)
    saved = _write_docx(spec, OUT_PATH)
    if saved != OUT_PATH:
        print(
            f"Saved to {saved} (could not overwrite {OUT_PATH} — file may be open in Word). "
            "Close sprint13_planning.docx and copy/rename, or rerun."
        )
    else:
        print(f"Wrote {saved}")
    return 0


if __name__ == "__main__":
    try:
        raise SystemExit(main())
    except GenerationError as e:
        print(str(e), file=sys.stderr)
        raise SystemExit(1)

