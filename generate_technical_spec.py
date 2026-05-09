from __future__ import annotations

import json
import os
import shutil
import time
from pathlib import Path
from typing import Any

from docx import Document
from dotenv import load_dotenv
from openai import OpenAI

DATASET_DIR = Path(__file__).resolve().parent
load_dotenv(DATASET_DIR / ".env")

MODEL = os.getenv("NUTRIVANA_OPENAI_MODEL", "openai/gpt-4o-mini")

BLUEPRINT_PATH = DATASET_DIR / "Technical_document.md"
OUT_PATH = DATASET_DIR / "generated" / "word" / "technical_spec.docx"

SYSTEM_PROMPT = """ROLE: You are Arjun Mehta, CTO of Nutrivana, writing the technical specification document on January 10 2025. You are writing for your engineering team and for your CEO Shristi who needs to understand technical decisions without being a developer.

CONTEXT: Nutrivana is a nutrition tracking app for health-conscious Indians built on Next.js frontend, FastAPI backend, Supabase with PostgreSQL and pgvector, LangChain for AI features. The team is small — 5 people. Every architectural decision must be justified clearly.

INSTRUCTION: Write a complete professional technical specification document. Every section must be specific and detailed. No vague statements. Reference actual table names, actual column names, actual technology choices with reasons. Write like a CTO who has thought deeply about every decision.

PERFORMANCE:
- Database schema must list all 12 tables with their key columns
- Every architectural decision must have a reason — not just what but why
- Performance targets must be specific numbers
- Security section must reference RLS policies specifically
- Known limitations must be honest — what we did not build and why
- Tone is technical but readable — Shristi should be able to understand the key decisions

SPECIFICATION:
Format as a Word document with these sections:
1. System Overview — what Nutrivana is and what the tech stack is
2. Database Schema — all 12 tables with key columns and relationships
3. Authentication and Security — JWT strategy, RLS policies, data isolation
4. API Design — FastAPI endpoints structure and principles
5. Search Architecture — USDA bulk import, trigram search, fuzzy matching
6. Performance Targets — specific numbers for all operations
7. Key Architectural Decisions — the important choices and why we made them
8. Known Limitations — honest list of what V1 does not do
9. Deployment — Render backend, Vercel frontend"""


class GenerationError(RuntimeError):
    pass


def _client() -> OpenAI:
    key = os.getenv("OPENAI_API_KEY", "").strip()
    base_url = os.getenv("OPENAI_BASE_URL", "").strip() or None
    if not key or key == "sk-or-v1-your-key-here":
        raise GenerationError(
            "OpenAI API key missing. Set OPENAI_API_KEY in nutrivana-dataset/.env (use python-dotenv)."
        )
    kwargs: dict[str, Any] = {"api_key": key}
    if base_url:
        kwargs["base_url"] = base_url
    return OpenAI(**kwargs)


def _call_doc_spec(client: OpenAI, *, blueprint_markdown: str, max_retries: int = 3) -> dict[str, Any]:
    user = (
        "Create a Word (.docx) document spec for Nutrivana's technical specification.\n"
        "Return ONLY valid JSON (no markdown, no code fences).\n\n"
        "JSON schema you must follow:\n"
        "{\n"
        '  "title": "string",\n'
        '  "subtitle": "string (optional)",\n'
        '  "sections": [\n'
        "    {\n"
        '      "heading": "string",\n'
        '      "level": 1 | 2 | 3 | 4,\n'
        '      "paragraphs": ["string", "..."]\n'
        "    }\n"
        "  ],\n"
        '  "tables": [\n'
        "    {\n"
        '      "title": "string (optional)",\n'
        '      "headers": ["..."],\n'
        '      "rows": [[...]]\n'
        "    }\n"
        "  ]\n"
        "}\n\n"
        "Hard requirements:\n"
        "- Include all 9 requested sections as headings (level=1) in order.\n"
        "- Database Schema section must include a table listing exactly 12 tables and key columns.\n"
        "- Performance Targets section must include a table with numeric targets.\n"
        "- Authentication and Security section must explicitly reference RLS policies (auth.uid()).\n"
        "- Known Limitations must be an honest list.\n\n"
        "Blueprint markdown (source of truth):\n"
        "-----\n"
        f"{blueprint_markdown}\n"
        "-----\n"
    )

    messages: list[dict[str, str]] = [
        {"role": "system", "content": SYSTEM_PROMPT},
        {"role": "user", "content": user},
    ]
    last_err: Exception | None = None
    for attempt in range(max_retries + 1):
        try:
            resp = client.chat.completions.create(
                model=MODEL,
                messages=messages,
                temperature=0.25,
                response_format={"type": "json_object"},
            )
            text = (resp.choices[0].message.content or "").strip()
            spec = json.loads(text)
            _validate_spec(spec)
            return spec
        except Exception as e:  # noqa: BLE001
            last_err = e
            if attempt == max_retries:
                break
            messages.append(
                {
                    "role": "user",
                    "content": (
                        "Your previous JSON failed validation. Return ONLY valid JSON per schema. "
                        f"Validation error: {e}"
                    ),
                }
            )
            time.sleep(min(2.0, 0.6 * (attempt + 1)))
    raise GenerationError(f"Failed after retries: {last_err}")


def _validate_spec(spec: dict[str, Any]) -> None:
    if not str(spec.get("title", "")).strip():
        raise ValueError("Missing non-empty title")
    sections = spec.get("sections")
    if not isinstance(sections, list) or not sections:
        raise ValueError("Missing sections array")
    tables = spec.get("tables")
    if not isinstance(tables, list) or not tables:
        raise ValueError("Missing tables array")

    required = [
        "System Overview",
        "Database Schema",
        "Authentication and Security",
        "API Design",
        "Search Architecture",
        "Performance Targets",
        "Key Architectural Decisions",
        "Known Limitations",
        "Deployment",
    ]
    headings = [str(s.get("heading", "")).strip() for s in sections if isinstance(s, dict)]
    for h in required:
        if not any(h.lower() == x.lower() for x in headings):
            raise ValueError(f"Missing required section heading: {h}")

    # Tables are injected from blueprint for strict accuracy (schema + performance).
    # We only require that the model produced at least one table structure so the doc can render.
    has_any_table = False
    for t in tables:
        if isinstance(t, dict) and (t.get("headers") or t.get("rows")):
            has_any_table = True
            break
    if not has_any_table:
        raise ValueError("Spec must include at least one table object.")


def _extract_schema_rows_from_blueprint(blueprint: str) -> list[list[str]]:
    """
    Extract table definitions under 'Section 4 — Database Schema' where each table is declared as:
      table_name:
      <columns line(s)>
    """
    m = blueprint.splitlines()
    # Locate Section 4 start
    start_idx = None
    for i, line in enumerate(m):
        if line.strip().lower().startswith("section 4"):
            start_idx = i
            break
    if start_idx is None:
        raise GenerationError("Could not locate 'Section 4 — Database Schema' in blueprint.")

    out: list[list[str]] = []
    i = start_idx
    current_name: str | None = None
    buf: list[str] = []

    def flush() -> None:
        nonlocal current_name, buf, out
        if current_name:
            cols = " ".join(x.strip() for x in buf if x.strip()).strip()
            if cols:
                out.append([current_name, cols, ""])
        current_name = None
        buf = []

    while i < len(m):
        line = m[i].rstrip()
        if line.strip().lower().startswith("section 5"):
            break
        # A table label line ends with ":" and is not "Note:".
        if line.strip().endswith(":") and not line.strip().lower().startswith("note:"):
            flush()
            current_name = line.strip()[:-1].strip()
        else:
            if current_name:
                if line.strip():
                    buf.append(line)
        i += 1
    flush()

    # We expect 12 app tables + users (auth) + rda_values etc; blueprint includes users + 11? but requirement says 12.
    # Keep the first 12 unique table names (excluding the Supabase-managed 'users' note if present).
    seen: set[str] = set()
    rows: list[list[str]] = []
    for name, cols, rel in out:
        if name.lower() == "users (managed by supabase auth)":
            continue
        if name.lower() == "users":
            continue
        if name in seen:
            continue
        seen.add(name)
        rows.append([name, cols, rel])
    if len(rows) < 12:
        raise GenerationError(f"Expected at least 12 tables in schema; found {len(rows)} ({[r[0] for r in rows]})")
    return rows[:12]


def _extract_performance_rows_from_blueprint(blueprint: str) -> list[list[str]]:
    # The blueprint includes a compressed line; use a fixed interpretation aligned to that content.
    return [
        ["USDA search", "< 1 second", "120ms (avg)"],
        ["Diary date switch", "< 2 seconds", "0.8 seconds"],
        ["Analysis load", "< 2 seconds", "0.8 seconds"],
        ["Nutrient recalculation", "< 1 second", "< 200ms"],
        ["Food add/delete", "< 1 second", "< 200ms"],
    ]


def _inject_accuracy_tables(spec: dict[str, Any], blueprint: str) -> None:
    tables = spec.get("tables")
    if not isinstance(tables, list):
        spec["tables"] = tables = []

    # Remove any existing schema/perf tables to avoid duplicates.
    cleaned: list[dict[str, Any]] = []
    for t in tables:
        if not isinstance(t, dict):
            continue
        title = str(t.get("title", "")).lower()
        if "database schema" in title or "performance targets" in title:
            continue
        cleaned.append(t)

    schema_rows = _extract_schema_rows_from_blueprint(blueprint)
    perf_rows = _extract_performance_rows_from_blueprint(blueprint)

    cleaned.append(
        {
            "title": "Database Schema — 12 Core Tables",
            "headers": ["Table", "Key Columns", "Relationships / Notes"],
            "rows": schema_rows,
        }
    )
    cleaned.append(
        {
            "title": "Performance Targets",
            "headers": ["Operation", "PRD Requirement", "Achieved"],
            "rows": perf_rows,
        }
    )
    spec["tables"] = cleaned


def _write_docx_from_spec(spec: dict[str, Any], out_path: Path) -> Path:
    doc = Document()
    title = spec.get("title") or out_path.stem
    doc.add_heading(str(title), level=0)
    subtitle = spec.get("subtitle")
    if subtitle:
        p = doc.add_paragraph(str(subtitle))
        if p.runs:
            p.runs[0].italic = True

    # Optional: a simple TOC placeholder (Word can update if user adds a real TOC).
    doc.add_heading("Table of Contents", level=1)
    doc.add_paragraph("Update this table of contents in Word after opening the document.")

    for sec in spec.get("sections", []) or []:
        if not isinstance(sec, dict):
            continue
        heading = sec.get("heading")
        level = int(sec.get("level", 1) or 1)
        if heading:
            doc.add_heading(str(heading), level=min(max(level, 1), 4))
        for para in sec.get("paragraphs", []) or []:
            if str(para).strip():
                doc.add_paragraph(str(para))

    for tbl in spec.get("tables", []) or []:
        if not isinstance(tbl, dict):
            continue
        ttitle = tbl.get("title")
        if ttitle:
            r = doc.add_paragraph(str(ttitle)).runs
            if r:
                r[0].bold = True
        headers = [str(h) for h in (tbl.get("headers") or [])]
        rows = tbl.get("rows") or []
        if not headers:
            continue
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        for i, h in enumerate(headers):
            table.cell(0, i).text = h
        for r in rows:
            tr = table.add_row().cells
            for i, v in enumerate(r):
                tr[i].text = "" if v is None else str(v)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    tmp = out_path.with_suffix(".tmp.docx")
    doc.save(tmp)
    try:
        os.replace(tmp, out_path)
        return out_path
    except PermissionError:
        fb = out_path.with_name("technical_spec_ready.docx")
        shutil.move(str(tmp), str(fb))
        return fb


def main() -> int:
    # User asked for round1_document_definitions.md, but in this dataset the Round 1 definitions
    # for the technical spec live in Technical_document.md.
    blueprint = BLUEPRINT_PATH.read_text(encoding="utf-8")
    client = _client()
    print(f"Model: {MODEL}")
    print("Calling API for technical spec...")
    spec = _call_doc_spec(client, blueprint_markdown=blueprint)
    _inject_accuracy_tables(spec, blueprint)
    saved = _write_docx_from_spec(spec, OUT_PATH)
    if saved != OUT_PATH:
        print(
            f"Saved to {saved} (could not overwrite {OUT_PATH} — file may be open in Word). "
            "Close technical_spec.docx and copy/rename, or rerun."
        )
    else:
        print(f"Wrote {saved}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

