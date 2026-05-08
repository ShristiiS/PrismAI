from __future__ import annotations

import json
import os
import re
import time
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from dotenv import load_dotenv
from docx import Document
from openai import OpenAI
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter


DATASET_DIR = Path(__file__).resolve().parent
load_dotenv(DATASET_DIR / ".env")
BLUEPRINT_DIR = DATASET_DIR
GENERATED_DIR = DATASET_DIR / "generated"
OUT_JIRA_DIR = GENERATED_DIR / "jira"
OUT_WORD_DIR = GENERATED_DIR / "word"
OUT_EXCEL_DIR = GENERATED_DIR / "excel"
OUT_EMAIL_DIR = GENERATED_DIR / "emails"

MODEL = os.getenv("NUTRIVANA_OPENAI_MODEL", "openai/gpt-4o-mini")


class GenerationError(RuntimeError):
    pass


def _ensure_dirs() -> None:
    OUT_JIRA_DIR.mkdir(parents=True, exist_ok=True)
    OUT_WORD_DIR.mkdir(parents=True, exist_ok=True)
    OUT_EXCEL_DIR.mkdir(parents=True, exist_ok=True)
    OUT_EMAIL_DIR.mkdir(parents=True, exist_ok=True)


def _slugify(name: str) -> str:
    s = name.strip().lower()
    s = re.sub(r"[^a-z0-9]+", "_", s)
    s = re.sub(r"_+", "_", s).strip("_")
    return s or "document"


def _read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def _extract_format_hint(markdown: str) -> str | None:
    m = re.search(r"(?:\*\*Format:\*\*|Format:)\s*([A-Za-z]+)\s*\(\.(docx|xlsx|txt)\)", markdown, re.I)
    if not m:
        return None
    kind = m.group(1).strip().lower()
    ext = m.group(2).strip().lower()
    if ext == "docx" or kind == "word":
        return "word"
    if ext == "xlsx" or kind == "excel":
        return "excel"
    if ext == "txt":
        return "txt"
    return None


def _is_sprint_ticket_blueprint(path: Path) -> int | None:
    name = path.name.lower()
    m = re.match(r"^sprint_(\d+)\.md$", name)
    if m:
        return int(m.group(1))
    m2 = re.match(r"^sprint(\d+)_ticket_definitions\.md$", name)
    if m2:
        return int(m2.group(1))
    return None


def _split_email_threads(markdown: str) -> list[tuple[str, str]]:
    parts = re.split(r"\n(?=## THREAD\s+—\s+)", "\n" + markdown.strip() + "\n")
    out: list[tuple[str, str]] = []
    for p in parts:
        p = p.strip()
        if not p:
            continue
        title_m = re.search(r"^## THREAD\s+—\s+(.+)$", p, re.M)
        title = title_m.group(1).strip() if title_m else "thread"
        out.append((title, p))
    return out


@dataclass(frozen=True)
class Blueprint:
    path: Path
    content: str


def _client() -> OpenAI:
    key = os.getenv("OPENAI_API_KEY", "").strip()
    base_url = os.getenv("OPENAI_BASE_URL", "").strip() or None

    if not key or key == "sk-or-v1-your-key-here":
        env_path = DATASET_DIR / ".env"
        print(
            f"[error] API key missing or placeholder. Set OPENAI_API_KEY in {env_path}. "
            f"Value length={len(key)}."
        )
        raise GenerationError(
            "OpenAI API key missing. Set OPENAI_API_KEY in nutrivana-dataset/.env (use python-dotenv)."
        )
    kwargs: dict[str, Any] = {"api_key": key}
    if base_url:
        kwargs["base_url"] = base_url
    return OpenAI(**kwargs)


def _call_gpt_json(client: OpenAI, *, system: str, user: str, schema_name: str, max_retries: int = 2) -> dict[str, Any]:
    last_err: Exception | None = None
    messages = [
        {"role": "system", "content": system},
        {"role": "user", "content": user},
    ]

    for attempt in range(max_retries + 1):
        try:
            resp = client.chat.completions.create(
                model=MODEL,
                messages=messages,
                temperature=0.2,
                response_format={"type": "json_object"},
            )
            text = (resp.choices[0].message.content or "").strip()
            return json.loads(text)
        except Exception as e:  # noqa: BLE001
            last_err = e
            messages.append(
                {
                    "role": "user",
                    "content": (
                        f"The previous answer did not parse as valid JSON for schema '{schema_name}'. "
                        "Return ONLY a valid JSON object, no markdown, no code fences, no commentary."
                    ),
                }
            )
            time.sleep(0.8 * (attempt + 1))

    raise GenerationError(f"Failed to get valid JSON from model after retries. Last error: {last_err}")


def _autosize_columns(ws) -> None:
    for col in range(1, ws.max_column + 1):
        max_len = 0
        for row in range(1, ws.max_row + 1):
            v = ws.cell(row=row, column=col).value
            if v is None:
                continue
            max_len = max(max_len, len(str(v)))
        ws.column_dimensions[get_column_letter(col)].width = min(max(10, max_len + 2), 60)


def _write_excel_from_spec(spec: dict[str, Any], out_path: Path) -> None:
    wb = Workbook()
    wb.remove(wb.active)

    for sheet in spec.get("sheets", []):
        name = str(sheet.get("name", "Sheet"))[:31]
        ws = wb.create_sheet(title=name)
        ws.freeze_panes = "A2"

        stype = (sheet.get("type") or "table").lower()
        if stype == "narrative":
            ws["A1"] = sheet.get("content", "")
            ws["A1"].alignment = Alignment(wrap_text=True, vertical="top")
            ws.row_dimensions[1].height = 400
            ws.column_dimensions["A"].width = 120
            continue

        headers = sheet.get("headers") or []
        rows = sheet.get("rows") or []

        if headers:
            ws.append([str(h) for h in headers])
            header_fill = PatternFill("solid", fgColor="1F4E79")
            header_font = Font(color="FFFFFF", bold=True)
            for c in range(1, len(headers) + 1):
                cell = ws.cell(row=1, column=c)
                cell.fill = header_fill
                cell.font = header_font
                cell.alignment = Alignment(wrap_text=True, vertical="top")

        for r in rows:
            ws.append(list(r))

        for row in ws.iter_rows(min_row=2, max_row=ws.max_row, min_col=1, max_col=ws.max_column):
            for cell in row:
                cell.alignment = Alignment(wrap_text=True, vertical="top")

        _autosize_columns(ws)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(out_path)


def _write_docx_from_spec(spec: dict[str, Any], out_path: Path) -> None:
    doc = Document()
    title = spec.get("title") or out_path.stem
    doc.add_heading(str(title), level=0)
    subtitle = spec.get("subtitle")
    if subtitle:
        p = doc.add_paragraph(str(subtitle))
        if p.runs:
            p.runs[0].italic = True

    for sec in spec.get("sections", []):
        heading = sec.get("heading")
        level = int(sec.get("level", 1))
        if heading:
            doc.add_heading(str(heading), level=min(max(level, 1), 4))
        for para in sec.get("paragraphs", []) or []:
            doc.add_paragraph(str(para))

    for tbl in spec.get("tables", []) or []:
        ttitle = tbl.get("title")
        if ttitle:
            r = doc.add_paragraph(str(ttitle)).runs
            if r:
                r[0].bold = True
        headers = [str(h) for h in (tbl.get("headers") or [])]
        rows = tbl.get("rows") or []
        if not headers:
            continue
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        for i, h in enumerate(headers):
            table.cell(0, i).text = h
        for r in rows:
            tr = table.add_row().cells
            for i, v in enumerate(r):
                tr[i].text = "" if v is None else str(v)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def _prompt_excel(blueprint: Blueprint) -> tuple[str, str]:
    system = (
        "You generate synthetic business spreadsheets from blueprint markdown.\n"
        "Return ONLY valid JSON (no markdown). Must match the requested sheet structure.\n"
        "Do not invent new metrics that contradict the blueprint. Keep names, dates, and numbers consistent.\n"
    )
    user = (
        "Create an Excel workbook spec from this blueprint.\n\n"
        "JSON schema you must follow:\n"
        "{\n"
        '  "workbook_title": "string",\n'
        '  "sheets": [\n'
        "    {\n"
        '      "name": "string",\n'
        '      "type": "table" | "narrative",\n'
        '      "headers": ["..."],\n'
        '      "rows": [[...]],\n'
        '      "content": "string"\n'
        "    }\n"
        "  ]\n"
        "}\n\n"
        "Rules:\n"
        "- For type=table: provide headers+rows; omit content.\n"
        "- For type=narrative: provide content; omit headers/rows.\n"
        "- Sheet names must be <= 31 chars.\n\n"
        f"Blueprint file: {blueprint.path.name}\n\n"
        "Blueprint markdown:\n"
        "-----\n"
        f"{blueprint.content}\n"
        "-----\n"
    )
    return system, user


def _prompt_word(blueprint: Blueprint) -> tuple[str, str]:
    system = (
        "You generate professional Word document content from blueprint markdown.\n"
        "Return ONLY valid JSON (no markdown). Use clear headings and realistic language.\n"
        "Do not contradict dates/numbers/names in the blueprint.\n"
    )
    user = (
        "Create a Word (.docx) document spec from this blueprint.\n\n"
        "JSON schema you must follow:\n"
        "{\n"
        '  "title": "string",\n'
        '  "subtitle": "string (optional)",\n'
        '  "sections": [\n'
        "    {\n"
        '      "heading": "string",\n'
        '      "level": 1 | 2 | 3 | 4,\n'
        '      "paragraphs": ["string", "..."]\n'
        "    }\n"
        "  ],\n"
        '  "tables": [\n'
        "    {\n"
        '      "title": "string (optional)",\n'
        '      "headers": ["..."],\n'
        '      "rows": [[...]]\n'
        "    }\n"
        "  ]\n"
        "}\n\n"
        f"Blueprint file: {blueprint.path.name}\n\n"
        "Blueprint markdown:\n"
        "-----\n"
        f"{blueprint.content}\n"
        "-----\n"
    )
    return system, user


def _prompt_jira_sprint(blueprint: Blueprint, sprint_number: int) -> tuple[str, str]:
    system = (
        "You convert Jira ticket blueprint markdown into a structured spreadsheet dataset.\n"
        "Return ONLY valid JSON. Keep fields consistent with the blueprint.\n"
    )
    user = (
        f"Create an Excel workbook spec for Sprint {sprint_number} Jira tickets.\n\n"
        "JSON schema you must follow:\n"
        "{\n"
        '  "workbook_title": "string",\n'
        '  "sheets": [\n'
        "    {\n"
        '      "name": "Tickets",\n'
        '      "type": "table",\n'
        '      "headers": ["Key","Title","Type","Priority","Status","Reporter","Assignee","Created","Resolved","Sprint","Tier","Participants","PRD Requirements","Conversation Blueprint","Unique Detail","Cross References"],\n'
        '      "rows": [[...]]\n'
        "    }\n"
        "  ]\n"
        "}\n\n"
        "Rules:\n"
        "- Put multi-line fields as single cell strings with newlines.\n"
        "- One row per ticket.\n\n"
        f"Blueprint file: {blueprint.path.name}\n\n"
        "Blueprint markdown:\n"
        "-----\n"
        f"{blueprint.content}\n"
        "-----\n"
    )
    return system, user


def _prompt_email_thread(thread_title: str, thread_markdown: str) -> tuple[str, str]:
    system = (
        "You generate realistic email threads from blueprint definitions.\n"
        "Return ONLY valid JSON. Output must be plain-text email thread content suitable for saving as .txt.\n"
        "Match participants, dates/times, and email count exactly.\n"
    )
    user = (
        "Generate the email thread as a single plain-text string.\n\n"
        "JSON schema you must follow:\n"
        "{\n"
        '  "thread_title": "string",\n'
        '  "output_txt": "string"\n'
        "}\n\n"
        "Formatting rules for output_txt:\n"
        "- Each email must have headers: From, To, Date, Subject.\n"
        "- Separate emails with a line of 72 hyphens.\n"
        "- Do NOT include markdown.\n\n"
        f"Thread title: {thread_title}\n\n"
        "Thread blueprint markdown:\n"
        "-----\n"
        f"{thread_markdown}\n"
        "-----\n"
    )
    return system, user


def _print_plan() -> None:
    print("Nutrivana document generator")
    print(f"- Dataset: {DATASET_DIR}")
    print(f"- Model: {MODEL}")
    print(f"- Output: {GENERATED_DIR}")


def _blueprint_paths() -> list[Path]:
    return sorted(BLUEPRINT_DIR.glob("*.md"), key=lambda p: p.name.lower())


def _already_done(path: Path) -> bool:
    return path.exists() and path.stat().st_size > 0


def main() -> int:
    _ensure_dirs()
    _print_plan()

    client = _client()
    blueprints = [Blueprint(path=p, content=_read_text(p)) for p in _blueprint_paths()]
    if not blueprints:
        print("No .md blueprints found.")
        return 1

    generated = 0
    skipped = 0

    for bp in blueprints:
        sprint_no = _is_sprint_ticket_blueprint(bp.path)
        if sprint_no is not None:
            out_path = OUT_JIRA_DIR / f"sprint_{sprint_no}.xlsx"
            if _already_done(out_path):
                print(f"[skip] Jira sprint {sprint_no}: {out_path.name}")
                skipped += 1
                continue
            system, user = _prompt_jira_sprint(bp, sprint_no)
            spec = _call_gpt_json(client, system=system, user=user, schema_name=f"jira_sprint_{sprint_no}")
            _write_excel_from_spec(spec, out_path)
            print(f"[ok]   Jira sprint {sprint_no}: {out_path.name}")
            generated += 1
            continue

        if bp.path.name.lower() == "email.md":
            threads = _split_email_threads(bp.content)
            for idx, (title, thread_md) in enumerate(threads, start=1):
                out_path = OUT_EMAIL_DIR / f"{idx:02d}_{_slugify(title)[:60]}.txt"
                if _already_done(out_path):
                    print(f"[skip] Email thread: {out_path.name}")
                    skipped += 1
                    continue
                system, user = _prompt_email_thread(title, thread_md)
                data = _call_gpt_json(client, system=system, user=user, schema_name=f"email_thread_{idx}")
                txt = str(data.get("output_txt", "")).strip()
                if not txt:
                    raise GenerationError(f"Empty email output for thread '{title}'")
                out_path.write_text(txt, encoding="utf-8")
                print(f"[ok]   Email thread: {out_path.name}")
                generated += 1
            continue

        fmt = _extract_format_hint(bp.content)
        if fmt == "excel":
            out_path = OUT_EXCEL_DIR / f"{bp.path.stem}.xlsx"
            if _already_done(out_path):
                print(f"[skip] Excel: {out_path.name}")
                skipped += 1
                continue
            system, user = _prompt_excel(bp)
            spec = _call_gpt_json(client, system=system, user=user, schema_name=f"excel_{bp.path.stem}")
            _write_excel_from_spec(spec, out_path)
            print(f"[ok]   Excel: {out_path.name}")
            generated += 1
            continue

        if fmt == "word":
            out_path = OUT_WORD_DIR / f"{bp.path.stem}.docx"
            if _already_done(out_path):
                print(f"[skip] Word: {out_path.name}")
                skipped += 1
                continue
            system, user = _prompt_word(bp)
            spec = _call_gpt_json(client, system=system, user=user, schema_name=f"word_{bp.path.stem}")
            _write_docx_from_spec(spec, out_path)
            print(f"[ok]   Word: {out_path.name}")
            generated += 1
            continue

        out_path = OUT_WORD_DIR / f"{bp.path.stem}.docx"
        if _already_done(out_path):
            print(f"[skip] Word(default): {out_path.name}")
            skipped += 1
            continue
        system, user = _prompt_word(bp)
        spec = _call_gpt_json(client, system=system, user=user, schema_name=f"word_default_{bp.path.stem}")
        _write_docx_from_spec(spec, out_path)
        print(f"[ok]   Word(default): {out_path.name}")
        generated += 1

    print()
    print(f"Done. Generated: {generated}, skipped: {skipped}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

